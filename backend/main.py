from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import PlainTextResponse
import io
from converter import cyrillic_to_tote, tote_to_cyrillic

app = FastAPI(title="KazScript API", version="1.0.0")

# CORS баптау
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "https://kazscript.vercel.app"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

MAX_FREE_SIZE = 500 * 1024  # 500 KB

@app.get("/")
def root():
    return {"message": "KazScript API жұмыс істеп тұр!", "version": "1.0.0"}

@app.get("/health")
def health():
    return {"status": "ok"}

@app.post("/convert/txt")
async def convert_txt(
    file: UploadFile = File(...),
    direction: str = "cyr2tote"
):
    # Файл форматын тексеру
    if not file.filename.endswith('.txt'):
        raise HTTPException(status_code=400, detail="Тек .txt файлдары қабылданады")

    # Файл көлемін тексеру
    contents = await file.read()
    if len(contents) > MAX_FREE_SIZE:
        raise HTTPException(status_code=400, detail="Файл көлемі 500 KB-тан аспауы керек")

    # Мәтінді оқу
    try:
        text = contents.decode('utf-8')
    except UnicodeDecodeError:
        try:
            text = contents.decode('cp1251')
        except:
            raise HTTPException(status_code=400, detail="Файл форматын оқу мүмкін емес")

    # Түрлендіру (алгоритм кейін қосылады)
    if direction == "cyr2tote":
        converted = cyrillic_to_tote(text)
    else:
        converted = tote_to_cyrillic(text)

    return PlainTextResponse(content=converted, media_type="text/plain; charset=utf-8")


from docx import Document
from docx.shared import Pt
import io

MAX_PREMIUM_SIZE = 10 * 1024 * 1024  # 10 MB

@app.post("/convert/docx")
async def convert_docx(
    file: UploadFile = File(...),
    direction: str = "cyr2tote"
):
    # Файл форматын тексеру
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Тек .docx файлдары қабылданады")

    # Файл көлемін тексеру
    contents = await file.read()
    if len(contents) > MAX_PREMIUM_SIZE:
        raise HTTPException(status_code=400, detail="Файл көлемі 10 MB-тан аспауы керек")

    # .docx файлын оқу
    doc = Document(io.BytesIO(contents))
    new_doc = Document()

    # Әр параграфты түрлендіру
    for para in doc.paragraphs:
        new_para = new_doc.add_paragraph()
        new_para.style = new_doc.styles['Normal']
        
        if para.text.strip():
            if direction == "cyr2tote":
                converted_text = cyrillic_to_tote(para.text)
            else:
                converted_text = tote_to_cyrillic(para.text)
            new_para.add_run(converted_text)
        else:
            new_para.add_run('')

    # Жаңа .docx файлын жіберу
    output = io.BytesIO()
    new_doc.save(output)
    output.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=converted.docx"}
    )